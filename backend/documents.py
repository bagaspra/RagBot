import os
import uuid
import logging
import asyncio
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import List, Optional, Dict

from fastapi import APIRouter, UploadFile, File, BackgroundTasks, HTTPException, Depends, Header
from pydantic import BaseModel, Field

from langchain_core.documents import Document
from backend.vector_store import get_vector_store, AppSettings, VectorStoreAdapter
from backend.ingest import _split_documents, _load_pdf, _load_txt

# Configure logging
logger = logging.getLogger(__name__)

router = APIRouter(prefix="/admin", tags=["admin"])
settings = AppSettings()

# --- Models ---

class DocumentStatus(str, Enum):
    PROCESSING = "processing"
    INDEXED    = "indexed"
    FAILED     = "failed"

class DocumentRecord(BaseModel):
    id: str
    filename: str
    file_type: str
    file_size_bytes: int
    uploaded_by: str
    uploaded_at: datetime
    status: DocumentStatus
    chunk_count: int = 0
    error_message: Optional[str] = None

class AdminStats(BaseModel):
    total_documents: int
    indexed_count: int
    processing_count: int
    failed_count: int
    total_size_bytes: int
    total_chunks: int
    qdrant_status: str
    groq_status: str

# --- In-memory Store & Persistence ---

# For production, this should be a real database (SQLite/Postgres)
# For now, we use a dictionary and rebuild from disk on startup
documents_store: Dict[str, DocumentRecord] = {}
DATA_DIR = Path("./data")
DATA_DIR.mkdir(exist_ok=True)

def _rebuild_store_from_disk():
    """Scan the data directory and rebuild the documents_store on startup."""
    global documents_store
    if not DATA_DIR.exists():
        return

    for file_path in DATA_DIR.glob("*"):
        if not file_path.is_file():
            continue
        
        # Expecting filenames like {uuid}_{original_filename}
        try:
            parts = file_path.name.split("_", 1)
            if len(parts) < 2:
                continue
            
            doc_id = parts[0]
            filename = parts[1]
            
            # Simple heuristic for rebuilding metadata
            stats = file_path.stat()
            documents_store[doc_id] = DocumentRecord(
                id=doc_id,
                filename=filename,
                file_type=file_path.suffix.lower().replace(".", ""),
                file_size_bytes=stats.st_size,
                uploaded_by="system", # Original uploader info lost on restart in memory-only mode
                uploaded_at=datetime.fromtimestamp(stats.st_ctime),
                status=DocumentStatus.INDEXED # Assume indexed if file exists on disk
            )
        except Exception as e:
            logger.error(f"Error rebuilding metadata for {file_path}: {e}")

_rebuild_store_from_disk()

# --- Middleware ---

async def verify_admin_key(x_admin_api_key: str = Header(None)):
    """Middleware to check for ADMIN_API_KEY in headers."""
    expected_key = AppSettings().ADMIN_API_KEY
    if not x_admin_api_key or x_admin_api_key != expected_key:
        raise HTTPException(status_code=403, detail="Forbidden: Invalid or missing Admin API Key")
    return x_admin_api_key

# --- Ingestion Helpers ---

async def _ingest_document_background(doc_id: str, file_path: Path):
    """Run ingestion in background and update status."""
    try:
        adapter = get_vector_store()
        
        # Load based on type
        suffix = file_path.suffix.lower()
        raw_docs = []
        
        if suffix == ".pdf":
            raw_docs = _load_pdf(file_path)
        elif suffix == ".txt":
            raw_docs = _load_txt(file_path)
        elif suffix == ".docx":
            # Use python-docx directly (no langchain_community dependency)
            import docx
            doc_obj = docx.Document(str(file_path))
            text = "\n".join(para.text for para in doc_obj.paragraphs if para.text.strip())
            raw_docs = [Document(page_content=text, metadata={"source": file_path.name})]
        elif suffix == ".xlsx":
            # Simple XLSX implementation using openpyxl
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            text = ""
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"
            raw_docs = [Document(page_content=text, metadata={"source": file_path.name})]
        
        if not raw_docs:
            raise ValueError("No content extracted from file")

        # Split and Embed
        chunks = _split_documents(raw_docs)
        adapter.add_documents(chunks)
        
        # Update store
        if doc_id in documents_store:
            documents_store[doc_id].status = DocumentStatus.INDEXED
            documents_store[doc_id].chunk_count = len(chunks)
            documents_store[doc_id].error_message = None
            
    except Exception as e:
        logger.error(f"Background ingestion failed for {doc_id}: {e}")
        if doc_id in documents_store:
            documents_store[doc_id].status = DocumentStatus.FAILED
            documents_store[doc_id].error_message = str(e)

# --- Endpoints ---

@router.get("/documents", response_model=List[DocumentRecord], dependencies=[Depends(verify_admin_key)])
async def get_documents(
    status: Optional[DocumentStatus] = None, 
    file_type: Optional[str] = None, 
    search: Optional[str] = None
):
    """List all documents with optional filtering."""
    results = list(documents_store.values())
    
    if status:
        results = [d for d in results if d.status == status]
    if file_type:
        results = [d for d in results if d.file_type == file_type]
    if search:
        search = search.lower()
        results = [d for d in results if search in d.filename.lower()]
        
    # Sort by date (newest first)
    results.sort(key=lambda x: x.uploaded_at, reverse=True)
    return results

@router.post("/documents/upload", response_model=List[DocumentRecord], dependencies=[Depends(verify_admin_key)])
async def upload_documents(
    background_tasks: BackgroundTasks,
    files: List[UploadFile] = File(...),
    uploaded_by: str = "admin"
):
    """Upload multiple documents and start background ingestion."""
    allowed_extensions = {".pdf", ".docx", ".txt", ".xlsx"}
    max_size = 50 * 1024 * 1024 # 50MB
    
    created_records = []
    
    for file in files:
        suffix = Path(file.filename).suffix.lower()
        if suffix not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"File type {suffix} not supported")
        
        # Read content to check size
        content = await file.read()
        if len(content) > max_size:
            raise HTTPException(status_code=413, detail=f"File {file.filename} is too large (max 50MB)")
        
        # Create record
        doc_id = str(uuid.uuid4())
        record = DocumentRecord(
            id=doc_id,
            filename=file.filename,
            file_type=suffix[1:], # remove dot
            file_size_bytes=len(content),
            uploaded_by=uploaded_by,
            uploaded_at=datetime.now(),
            status=DocumentStatus.PROCESSING
        )
        
        # Save to disk
        file_path = DATA_DIR / f"{doc_id}_{file.filename}"
        with open(file_path, "wb") as f:
            f.write(content)
            
        documents_store[doc_id] = record
        created_records.append(record)
        
        # Trigger background task
        background_tasks.add_task(_ingest_document_background, doc_id, file_path)
        
    return created_records

@router.delete("/documents/{document_id}", dependencies=[Depends(verify_admin_key)])
async def delete_document(document_id: str):
    """Delete a document, its file, and its vectors from Qdrant."""
    if document_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    
    record = documents_store[document_id]
    file_path = DATA_DIR / f"{document_id}_{record.filename}"
    
    # 1. Delete vectors from Qdrant
    try:
        from qdrant_client import models
        adapter = get_vector_store()
        # QdrantVectorStore handles deletion by metadata filter
        # But we might need to access the client directly for precise control
        q_client = adapter._client
        q_client.delete(
            collection_name=settings.QDRANT_COLLECTION_NAME,
            points_selector=models.Filter(
                must=[
                    models.FieldCondition(
                        key="metadata.source", # langchain_qdrant puts source in metadata
                        match=models.MatchValue(value=record.filename)
                    )
                ]
            )
        )
    except Exception as e:
        logger.warning(f"Failed to delete vectors for {record.filename}: {e}")

    # 2. Delete file
    if file_path.exists():
        file_path.unlink()
        
    # 3. Remove from store
    del documents_store[document_id]
    
    return {"deleted": True}

@router.post("/documents/{document_id}/reindex", response_model=DocumentRecord, dependencies=[Depends(verify_admin_key)])
async def reindex_document(document_id: str, background_tasks: BackgroundTasks):
    """Re-index an existing document."""
    if document_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    
    record = documents_store[document_id]
    file_path = DATA_DIR / f"{document_id}_{record.filename}"
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Source file not found on disk")
        
    # 1. Delete old vectors
    try:
        from qdrant_client import models
        adapter = get_vector_store()
        q_client = adapter._client
        q_client.delete(
            collection_name=settings.QDRANT_COLLECTION_NAME,
            points_selector=models.Filter(
                must=[
                    models.FieldCondition(
                        key="metadata.source",
                        match=models.MatchValue(value=record.filename)
                    )
                ]
            )
        )
    except Exception as e:
        logger.warning(f"Failed to delete old vectors during reindex for {record.filename}: {e}")

    # 2. Reset status
    record.status = DocumentStatus.PROCESSING
    record.chunk_count = 0
    record.error_message = None
    
    # 3. Start background task
    background_tasks.add_task(_ingest_document_background, document_id, file_path)
    
    return record

@router.get("/stats", response_model=AdminStats, dependencies=[Depends(verify_admin_key)])
async def get_admin_stats():
    """Get system-wide statistics."""
    records = list(documents_store.values())
    
    # Qdrant status
    qdrant_status = "online"
    total_chunks = 0
    try:
        adapter = get_vector_store()
        info = adapter._client.get_collection(settings.QDRANT_COLLECTION_NAME)
        total_chunks = info.points_count
    except Exception:
        qdrant_status = "offline"

    # Groq status
    groq_status = "online"
    # Basic health check for LLM provider could go here
    
    return AdminStats(
        total_documents=len(records),
        indexed_count=len([r for r in records if r.status == DocumentStatus.INDEXED]),
        processing_count=len([r for r in records if r.status == DocumentStatus.PROCESSING]),
        failed_count=len([r for r in records if r.status == DocumentStatus.FAILED]),
        total_size_bytes=sum(r.file_size_bytes for r in records),
        total_chunks=total_chunks,
        qdrant_status=qdrant_status,
        groq_status=groq_status
    )
